"""
Parse a Respondus-style DOCX assessment file into a Chapter + Question model.

Expected DOCX paragraph format (all paragraphs use 'Normal Text' style):

  Book Title Line           — skipped
  Chapter N   Title line    — sets chapter title
  N.N  Section heading      — skipped
  N) Question stem          — new question
  A) Choice text            — MC choice (A–D)
  ...
  Answer:  X                — correct answer (letter, TRUE/FALSE, or essay text)
  (continuation lines)      — multi-line essay answer, until Diff line
  Diff: N    Type: XX       — metadata; sets question type (MC/TF/SA/ES)
  Objective:  N-N           — skipped
"""

import html as _html_esc
import io
import re

from docx import Document

from .models import Chapter, Choice, Question

# Map Respondus type codes → Moodle question types
_TYPE_MAP = {
    "MC": "multichoice",
    "TF": "truefalse",
    "SA": "essay",
    "ES": "essay",
}

# Regex patterns
_RE_CHAPTER = re.compile(r"^Chapter\s+(\d+)\s{2,}(.+)", re.IGNORECASE)
_RE_SECTION = re.compile(r"^\d+\.\d+\s{2,}")
_RE_QUESTION = re.compile(r"^(\d+)\)\s+(.+)")
_RE_CHOICE = re.compile(r"^([A-D])\)\s+(.+)")
# The answer may be empty on the first line, with the actual essay answer
# beginning in one or more continuation paragraphs below it.
_RE_ANSWER = re.compile(r"^Answer:\s*(.*)", re.IGNORECASE)
_RE_DIFF = re.compile(r"^Diff:\s*\d+\s+Type:\s*(\w+)", re.IGNORECASE)
_RE_OBJECTIVE = re.compile(r"^Objective:\s+", re.IGNORECASE)


def parse_chapter(file_bytes: bytes) -> Chapter:
    doc = Document(io.BytesIO(file_bytes))

    # Pearson-style test banks use [Q1] markers and unlabeled answer
    # paragraphs. Detect that format before falling back to the Respondus
    # parser below.
    if any(re.fullmatch(r"\[Q\d+\]", para.text.strip(), re.IGNORECASE) for para in doc.paragraphs):
        return _parse_bracketed_test_bank(doc)

    chapter = Chapter(title="Untitled Chapter", number=0)
    current_q: Question | None = None
    in_answer = False
    global_num = 0
    first_para_seen_chapter = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Preserve blank paragraphs inside multi-paragraph essay answers.
            # The following non-empty continuation adds another <br>, yielding
            # the <br><br> spacing Blackboard renders as a paragraph gap.
            if in_answer and current_q is not None and current_q.model_answer:
                current_q.model_answer += "<br>"
            in_answer = False if _answer_ended(current_q, in_answer) else in_answer
            continue

        # ── Chapter title ────────────────────────────────────────────
        m = _RE_CHAPTER.match(text)
        if m and not first_para_seen_chapter:
            first_para_seen_chapter = True
            num = int(m.group(1))
            rest = m.group(2).strip()
            # Normalise multiple internal spaces in title
            rest = re.sub(r"\s{2,}", " ", rest)
            chapter.number = num
            chapter.title = f"Chapter {num:02d} {rest}"
            continue

        # ── Section heading (e.g. "1.2  Define …") ───────────────────
        if _RE_SECTION.match(text):
            continue

        # ── Diff / Type metadata line ─────────────────────────────────
        m = _RE_DIFF.match(text)
        if m:
            if current_q is not None:
                code = m.group(1).upper()
                current_q.q_type = _TYPE_MAP.get(code, "essay")
                _finalise_answer(current_q)
            in_answer = False
            continue

        # ── Objective metadata line ───────────────────────────────────
        if _RE_OBJECTIVE.match(text):
            continue

        # ── Answer line ───────────────────────────────────────────────
        m = _RE_ANSWER.match(text)
        if m and current_q is not None:
            answer_text = m.group(1).strip()
            current_q.correct_letter = answer_text  # raw; interpreted later
            # Build HTML model answer, stripping the "Answer:  " prefix
            html_full = _para_to_html(para)
            current_q.model_answer = re.sub(r"^Answer:\s+", "", html_full, flags=re.IGNORECASE).strip()
            in_answer = True
            continue

        # ── Essay answer continuation ─────────────────────────────────
        if in_answer and current_q is not None:
            # Stop continuation if a new question starts
            if _RE_QUESTION.match(text):
                in_answer = False
                # fall through to question handling below
            else:
                current_q.model_answer += "<br>" + _para_to_html(para)
                continue

        # ── New question ──────────────────────────────────────────────
        m = _RE_QUESTION.match(text)
        if m:
            global_num += 1
            # Extract HTML from the paragraph, strip the leading "N) " prefix
            html_full = _para_to_html(para)
            prefix = m.group(1) + ") "
            stem_html = html_full[len(prefix):] if html_full.startswith(prefix) else html_full
            current_q = Question(
                number=global_num,
                stem=stem_html.strip(),
                q_type="essay",  # overwritten by Diff line
            )
            chapter.questions.append(current_q)
            in_answer = False
            continue

        # ── MC choice (A–D) ───────────────────────────────────────────
        m = _RE_CHOICE.match(text)
        if m and current_q is not None and not in_answer:
            html_full = _para_to_html(para)
            prefix = m.group(1) + ") "
            choice_html = html_full[len(prefix):] if html_full.startswith(prefix) else html_full
            current_q.choices.append(Choice(letter=m.group(1), text=choice_html.strip()))
            continue

        # Anything else: skip (book title line, stray text, etc.)

    return chapter


def _parse_bracketed_test_bank(doc: Document) -> Chapter:
    """Parse test banks with [Qn] markers, four plain choices, and markers."""
    paragraphs = [para for para in doc.paragraphs if para.text.strip()]
    question_marker = re.compile(r"^\[Q\d+\]$", re.IGNORECASE)
    chapter_title = "Untitled Chapter"
    chapter_number = 0

    for para in paragraphs:
        match = re.match(r"^Chapter\s+(\d+)\s*:\s*(.+)$", para.text.strip(), re.IGNORECASE)
        if match:
            chapter_number = int(match.group(1))
            chapter_title = f"Chapter {chapter_number:02d} {re.sub(r'\s+', ' ', match.group(2).strip())}"
            break

    chapter = Chapter(title=chapter_title, number=chapter_number)
    index = 0
    while index < len(paragraphs):
        if not question_marker.fullmatch(paragraphs[index].text.strip()):
            index += 1
            continue

        index += 1
        if index >= len(paragraphs):
            break
        stem_para = paragraphs[index]
        index += 1
        # Application/model-response sections also use [Qn] markers, but do
        # not contain four plain answer choices. Do not turn those sections
        # into malformed multiple-choice items.
        if any(paragraphs[index + offset].text.strip().startswith("[") for offset in range(min(4, len(paragraphs) - index))):
            continue
        choices = []
        for letter in "ABCD":
            if index >= len(paragraphs) or question_marker.fullmatch(paragraphs[index].text.strip()):
                break
            choice_para = paragraphs[index]
            index += 1
            is_correct = bool(re.search(r"\s*\(correct\)\s*$", choice_para.text, re.IGNORECASE))
            # Strip the marker after converting to HTML as DOCX formatting can
            # wrap it in tags (for example, <b>(correct)</b>).
            choice_html = re.sub(r"\s*\(correct\)\s*", " ", _para_to_html(choice_para), flags=re.IGNORECASE).strip()
            choice_html = re.sub(r"<(b|i|u)>\s*</\1>", "", choice_html, flags=re.IGNORECASE)
            choices.append(Choice(letter=letter, text=choice_html))

        if len(choices) != 4:
            continue

        correct = next((choice.letter for choice, para in zip(choices, paragraphs[index - 4:index])
                        if re.search(r"\s*\(correct\)\s*$", para.text, re.IGNORECASE)), "")
        chapter.questions.append(Question(
            number=len(chapter.questions) + 1,
            stem=_para_to_html(stem_para).strip(),
            q_type="multichoice",
            choices=choices,
            correct_letter=correct,
        ))

    return chapter


# ─── helpers ──────────────────────────────────────────────────────────────────

def _answer_ended(q: Question | None, in_answer: bool) -> bool:
    """Blank line doesn't necessarily end multi-line essay answers."""
    return False  # we keep collecting until Diff line


def _para_to_html(para) -> str:
    """
    Convert a DOCX paragraph to an HTML string, preserving inline formatting
    (bold, italic, underline) and paragraph-level page-break-before style.
    Text characters <, >, and & are HTML-escaped so the output is valid HTML.
    """
    parts = []
    for run in para.runs:
        if not run.text:
            continue
        segment = _html_esc.escape(run.text)
        if run.bold and run.italic:
            segment = f"<b><i>{segment}</i></b>"
        elif run.bold:
            segment = f"<b>{segment}</b>"
        elif run.italic:
            segment = f"<i>{segment}</i>"
        if run.underline:
            segment = f"<u>{segment}</u>"
        parts.append(segment)

    content = "".join(parts)

    # Wrap in <p> if this paragraph has page-break-before styling
    try:
        if para.paragraph_format.page_break_before:
            content = f'<p style="page-break-before:always">{content}</p>'
    except Exception:
        pass

    return content


def _finalise_answer(q: Question) -> None:
    """
    Normalise the correct_letter field once the type is known.

    For MC  : correct_letter should be uppercase single letter ('A'..'D').
    For TF  : correct_letter should be 'TRUE' or 'FALSE'.
    For essay: correct_letter is irrelevant; model_answer holds the text.
    """
    if q.q_type == "multichoice":
        q.correct_letter = q.correct_letter.strip().upper()
    elif q.q_type == "truefalse":
        q.correct_letter = q.correct_letter.strip().upper()
    else:
        # essay / SA — model_answer is already set; strip leading/trailing whitespace
        q.model_answer = q.model_answer.strip()
